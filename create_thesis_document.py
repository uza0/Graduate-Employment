"""
Script to generate the JoinWork graduation thesis document
Creates a Word document with proper formatting
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_run_font(run, font_name='Times New Roman', size=12):
    """Set font properties for a run"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_with_ltr(doc, text, level=1):
    """Add a heading with LTR alignment"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for paragraph in doc.paragraphs:
        if paragraph.text == text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_ltr(doc, text, style=None):
    """Add a paragraph with LTR alignment"""
    if style:
        para = doc.add_paragraph(text, style=style)
    else:
        para = doc.add_paragraph(text)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return para

def create_thesis_document():
    """Create the graduation thesis document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Title Page
    title = doc.add_heading('JoinWork: A Digital Platform for Connecting Graduates with Employment Opportunities', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph('Graduation Thesis')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('Submitted to the College of Computer Science and Mathematics')
    doc.add_paragraph('University of Baghdad')
    doc.add_paragraph('In Partial Fulfillment of the Requirements for the Degree of Bachelor of Science in Computer Science')
    
    doc.add_page_break()
    
    # CHAPTER 1: INTRODUCTION
    add_heading_with_ltr(doc, 'CHAPTER 1: INTRODUCTION', 1)
    
    add_heading_with_ltr(doc, '1.1 General Introduction', 2)
    add_paragraph_with_ltr(doc, 
        'Unemployment among graduates represents a significant challenge in many countries, including Iraq. '
        'Many graduates struggle to find employment opportunities that match their qualifications and skills, '
        'while employers face difficulty in identifying and reaching qualified candidates. Traditional methods '
        'of job searching and recruitment, such as newspaper advertisements, informal networks, and unorganized '
        'social media platforms, have proven to be inefficient and unreliable.')
    add_paragraph_with_ltr(doc,
        'The digital transformation era has introduced innovative solutions to bridge this gap. Web-based platforms '
        'can centralize job listings, streamline application processes, and facilitate better matching between '
        'graduates and employers. JoinWork addresses this critical need by providing a centralized digital platform '
        'that connects graduates with employment opportunities while supporting companies in their recruitment efforts '
        'and enabling the Ministry to monitor employment trends.')
    
    add_heading_with_ltr(doc, '1.2 Problem Statement', 2)
    add_paragraph_with_ltr(doc,
        'The research problem is defined as the absence of a unified, centralized digital platform that effectively '
        'connects graduates with the job market. The specific issues addressed by this research include:')
    
    para = add_paragraph_with_ltr(doc, '1. Lack of Centralized Platform: No single system that organizes job opportunities and graduate profiles in one accessible location.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '2. Difficulty in CV Creation: Many graduates lack the skills or tools to create professional CVs that meet employer standards.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '3. Scattered Job Opportunities: Job postings are distributed across multiple channels (social media, newspapers, word-of-mouth), making it difficult for graduates to find relevant opportunities.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '4. Unreliable Information: Lack of verification mechanisms leads to unreliable job postings and candidate information.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '5. Absence of Organized Database: No systematic database of graduate profiles, skills, and qualifications that employers can efficiently search.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '6. Gap Between Academia and Industry: Limited connection between academic achievements and practical job market requirements.')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '1.3 Project Objectives', 2)
    add_paragraph_with_ltr(doc, 'The primary objectives of JoinWork are:')
    
    para = add_paragraph_with_ltr(doc, '1. Create a Digital Platform: Develop a web-based system that serves as a bridge between graduates and employers.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '2. Facilitate CV Creation: Provide an automated CV generation system that produces professional, well-formatted CVs based on user profile data.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '3. Enable Job Posting: Allow companies to post job opportunities with detailed requirements and specifications.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '4. Streamline Application Process: Simplify the job application process for graduates and application management for companies.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '5. Reduce the Gap: Bridge the gap between academic education and practical job market needs by providing relevant job opportunities and training workshops.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '6. Support Recent Graduates: Specifically target and support newly graduated students who are entering the job market for the first time.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '7. Provide Analytics: Offer the Ministry and universities insights into employment trends, graduate skills, and market demands.')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '1.4 General Research Design', 2)
    add_paragraph_with_ltr(doc,
        'This thesis is structured into five main chapters, each serving a specific purpose in presenting the research:')
    
    para = add_paragraph_with_ltr(doc, '• Chapter 1: Introduction - Presents the problem, objectives, and related studies. It establishes the context and motivation for the research.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '• Chapter 2: Theoretical and Technical Background - Covers the theoretical foundations, technologies, and tools used in the development of JoinWork, including web technologies, database systems, and authentication mechanisms.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '• Chapter 3: Algorithms and Flowcharts - Describes the system\'s algorithms and flowcharts for major processes, prepared by the researcher. This chapter explains the logical flow of operations without code implementation.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '• Chapter 4: Implementation and Results - Demonstrates the practical implementation of JoinWork, including interface screenshots, system functionality, and discussion of results.')
    para.style = 'List Bullet'
    
    para = add_paragraph_with_ltr(doc, '• Chapter 5: Conclusions and Future Work - Summarizes achievements, evaluates objective fulfillment, and proposes future enhancements and extensions.')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '1.5 Related and Previous Studies', 2)
    
    add_heading_with_ltr(doc, '1.5.1 LinkedIn', 3)
    add_paragraph_with_ltr(doc,
        'LinkedIn is a professional networking platform that connects professionals worldwide. While it offers job postings, '
        'networking, and profile management, it has certain limitations when it comes to serving recent graduates.')
    add_paragraph_with_ltr(doc, 'Strengths:')
    para = add_paragraph_with_ltr(doc, '• Global reach and extensive user base')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Advanced networking features')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Professional profile management')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Industry-specific groups')
    para.style = 'List Bullet'
    
    add_paragraph_with_ltr(doc, 'Weaknesses:')
    para = add_paragraph_with_ltr(doc, '• Not specifically designed for recent graduates')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Complex interface that may overwhelm new users')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Limited focus on local job markets')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Premium features require payment')
    para.style = 'List Bullet'
    
    add_paragraph_with_ltr(doc, 'JoinWork Differentiation:')
    para = add_paragraph_with_ltr(doc, '• Specifically targets graduates, especially recent ones')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Simplified interface for ease of use')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Focus on local (Iraqi) job market')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Free access to all features')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '1.5.2 Bayt.com', 3)
    add_paragraph_with_ltr(doc,
        'Bayt.com is a leading job portal in the Middle East, focusing on job listings and recruitment.')
    add_paragraph_with_ltr(doc, 'Strengths:')
    para = add_paragraph_with_ltr(doc, '• Extensive job listings')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Regional focus on Middle East')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Established reputation')
    para.style = 'List Bullet'
    
    add_paragraph_with_ltr(doc, 'Weaknesses:')
    para = add_paragraph_with_ltr(doc, '• Limited CV building tools')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Less emphasis on graduate-specific features')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Complex application processes')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Not tailored for university graduates')
    para.style = 'List Bullet'
    
    add_paragraph_with_ltr(doc, 'JoinWork Differentiation:')
    para = add_paragraph_with_ltr(doc, '• Automated CV generation with multiple templates')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Graduate-focused features (GPA, university, major tracking)')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Simplified application process')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Integration with university data')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '1.5.3 Local Employment Platforms', 3)
    add_paragraph_with_ltr(doc,
        'Various local platforms exist but often lack centralized graduate database, automated CV generation, '
        'Ministry/university integration, and bilingual support (Arabic/English).')
    
    add_heading_with_ltr(doc, '1.5.4 Summary of Competitive Advantages', 3)
    add_paragraph_with_ltr(doc, 'JoinWork distinguishes itself through:')
    para = add_paragraph_with_ltr(doc, '1. Graduate-Centric Design: Specifically designed for university graduates')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. Automated CV Generation: Multiple professional templates with one-click generation')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Local Focus: Tailored for the Iraqi job market')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. Ministry Integration: Special portal for government monitoring and analytics')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. Bilingual Support: Full Arabic and English interface')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. Free Access: No premium features or payment requirements')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. Simplified Interface: User-friendly design for graduates with varying technical skills')
    para.style = 'List Number'
    
    doc.add_page_break()
    
    # CHAPTER 2: THEORETICAL AND TECHNICAL BACKGROUND
    add_heading_with_ltr(doc, 'CHAPTER 2: THEORETICAL AND TECHNICAL BACKGROUND', 1)
    
    add_heading_with_ltr(doc, '2.1 Chapter Introduction', 2)
    add_paragraph_with_ltr(doc,
        'Before implementing any system, it is essential to establish a solid theoretical foundation. This chapter '
        'presents the theoretical concepts, technologies, and architectural patterns used in JoinWork, providing the '
        'academic and technical basis for the system\'s design and development.')
    
    add_heading_with_ltr(doc, '2.2 Adopted Theories and Concepts', 2)
    
    add_heading_with_ltr(doc, '2.2.1 Web-Based Systems', 3)
    add_paragraph_with_ltr(doc,
        'Web-based systems are applications that are accessed through web browsers, eliminating the need for software '
        'installation on client devices. JoinWork follows this model, enabling access from any device with internet connectivity.')
    add_paragraph_with_ltr(doc, 'Advantages:')
    para = add_paragraph_with_ltr(doc, '• Cross-platform compatibility')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Centralized updates')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• No installation required')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Easy maintenance')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '2.2.2 Client-Server Architecture', 3)
    add_paragraph_with_ltr(doc,
        'JoinWork employs a client-server architecture where the client (frontend) communicates with the server (backend) '
        'through HTTP requests. The frontend, built with HTML, CSS, and JavaScript, runs in the user\'s browser, while '
        'the backend, built with Python Flask, handles business logic, authentication, and data management.')
    add_paragraph_with_ltr(doc,
        'This separation of concerns improves scalability, maintainability, and security.')
    
    add_heading_with_ltr(doc, '2.2.3 Database Management Systems', 3)
    add_paragraph_with_ltr(doc,
        'Data persistence in JoinWork is managed through JSON file storage, providing structured data organization, '
        'data integrity, efficient retrieval, and scalability potential.')
    
    add_heading_with_ltr(doc, '2.2.4 Authentication and Authorization', 3)
    add_paragraph_with_ltr(doc,
        'The system implements authentication using JWT (JSON Web Tokens) for user verification and role-based access '
        'control (Graduate, Company, Ministry) for authorization. Password security is ensured through SHA-256 hashing.')
    
    add_heading_with_ltr(doc, '2.2.5 RESTful API Architecture', 3)
    add_paragraph_with_ltr(doc,
        'JoinWork follows REST principles with stateless communication, standard HTTP methods (GET, POST, PUT, DELETE), '
        'JSON data format, and resource-based URLs.')
    
    add_heading_with_ltr(doc, '2.3 Languages and Technologies Used', 2)
    
    add_heading_with_ltr(doc, '2.3.1 Frontend Technologies', 3)
    
    add_heading_with_ltr(doc, 'HTML (HyperText Markup Language)', 4)
    add_paragraph_with_ltr(doc, '• Creator: Tim Berners-Lee (1991)')
    add_paragraph_with_ltr(doc, '• Purpose: Structure and content of web pages')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: Page structure, forms, semantic elements')
    
    add_heading_with_ltr(doc, 'CSS (Cascading Style Sheets)', 4)
    add_paragraph_with_ltr(doc, '• Creator: Håkon Wium Lie (1994)')
    add_paragraph_with_ltr(doc, '• Purpose: Styling and layout')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: Responsive design, theme system, component styling')
    
    add_heading_with_ltr(doc, 'JavaScript', 4)
    add_paragraph_with_ltr(doc, '• Creator: Brendan Eich (1995)')
    add_paragraph_with_ltr(doc, '• Purpose: Client-side interactivity')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: Dynamic content, API calls, CV generation, form validation')
    
    add_heading_with_ltr(doc, '2.3.2 Backend Technologies', 3)
    
    add_heading_with_ltr(doc, 'Python', 4)
    add_paragraph_with_ltr(doc, '• Creator: Guido van Rossum (1991)')
    add_paragraph_with_ltr(doc, '• Purpose: Server-side logic')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: Business logic, API endpoints, data processing')
    
    add_heading_with_ltr(doc, 'Flask', 4)
    add_paragraph_with_ltr(doc, '• Creator: Armin Ronacher (2010)')
    add_paragraph_with_ltr(doc, '• Purpose: Web framework')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: RESTful API, routing, request handling')
    
    add_heading_with_ltr(doc, 'Flask-CORS', 4)
    add_paragraph_with_ltr(doc, '• Purpose: Cross-Origin Resource Sharing')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: Enabling frontend-backend communication')
    
    add_heading_with_ltr(doc, '2.3.3 Data Storage', 3)
    
    add_heading_with_ltr(doc, 'JSON (JavaScript Object Notation)', 4)
    add_paragraph_with_ltr(doc, '• Creator: Douglas Crockford (2001)')
    add_paragraph_with_ltr(doc, '• Purpose: Data interchange format')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: Persistent data storage, API communication')
    
    add_heading_with_ltr(doc, '2.3.4 Security Technologies', 3)
    
    add_heading_with_ltr(doc, 'JWT (JSON Web Tokens)', 4)
    add_paragraph_with_ltr(doc, '• Standard: RFC 7519')
    add_paragraph_with_ltr(doc, '• Purpose: Secure token-based authentication')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: User authentication and session management')
    
    add_heading_with_ltr(doc, 'SHA-256 (Secure Hash Algorithm 256-bit)', 4)
    add_paragraph_with_ltr(doc, '• Purpose: Password hashing')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: Secure password storage')
    
    add_heading_with_ltr(doc, '2.3.5 Additional Libraries', 3)
    
    add_heading_with_ltr(doc, 'html2pdf / jsPDF / html2canvas', 4)
    add_paragraph_with_ltr(doc, '• Purpose: PDF generation from HTML')
    add_paragraph_with_ltr(doc, '• Usage in JoinWork: CV export functionality')
    
    add_heading_with_ltr(doc, '2.4 Reasons for Technology Selection', 2)
    
    add_heading_with_ltr(doc, '2.4.1 Performance', 3)
    add_paragraph_with_ltr(doc,
        'Python Flask provides lightweight, fast API development. JSON storage offers fast read/write operations '
        'for moderate data volumes. JavaScript enables efficient client-side processing.')
    
    add_heading_with_ltr(doc, '2.4.2 Flexibility', 3)
    add_paragraph_with_ltr(doc,
        'The modular architecture allows easy extension and modification. The technology stack is well-documented '
        'and widely supported. The API design allows future integration with mobile applications or other systems.')
    
    add_heading_with_ltr(doc, '2.4.3 Maintainability', 3)
    add_paragraph_with_ltr(doc,
        'Clean code structure with organized frontend and backend separation. Standard technologies make it easy '
        'for developers to understand and maintain. Extensive community resources provide documentation and support.')
    
    add_heading_with_ltr(doc, '2.4.4 Suitability for Graduation Project', 3)
    add_paragraph_with_ltr(doc,
        'The technologies provide excellent learning value covering full-stack development concepts. The complexity '
        'level is appropriate for a graduation project. The system addresses a real-world problem. The architecture '
        'has scalability potential for future production deployment.')
    
    add_heading_with_ltr(doc, '2.4.5 Cost-Effectiveness', 3)
    add_paragraph_with_ltr(doc,
        'All technologies are open source and free. No commercial licenses are required. The system can be deployed '
        'on various platforms without additional costs.')
    
    doc.add_page_break()
    
    # CHAPTER 3: ALGORITHMS AND FLOWCHARTS
    add_heading_with_ltr(doc, 'CHAPTER 3: ALGORITHMS AND FLOWCHARTS', 1)
    
    add_heading_with_ltr(doc, '3.1 Chapter Introduction', 2)
    add_paragraph_with_ltr(doc,
        'This chapter presents the algorithms and flowcharts developed by the researcher for JoinWork\'s core processes. '
        'These represent the logical flow of operations without code implementation, focusing on the step-by-step procedures '
        'that govern system behavior.')
    
    add_heading_with_ltr(doc, '3.2 System Algorithms (Prepared by the Researcher)', 2)
    
    add_heading_with_ltr(doc, '3.2.1 User Registration Algorithm', 3)
    add_paragraph_with_ltr(doc, 'Purpose: Register new users (Graduates, Companies, or Ministry) in the system.')
    add_paragraph_with_ltr(doc, 'Steps:')
    para = add_paragraph_with_ltr(doc, '1. User accesses the registration page')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. User selects role (Graduate, Company, or Ministry)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. User fills required fields (name, email, password, role-specific data)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. System validates input (email format, password strength, required fields)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. System checks if email already exists')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. If email exists, display error message and return to step 3')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. If email is unique, hash the password using SHA-256')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '8. Generate unique user_id')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '9. Create user record with hashed password')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '10. If role is Graduate, create graduate profile with provided data')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '11. If role is Company, create company profile with company details')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '12. Save all data to persistent storage')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '13. Generate JWT token for authenticated session')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '14. Return success response with token and user information')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '15. Redirect user to appropriate dashboard')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '3.2.2 User Login Algorithm', 3)
    add_paragraph_with_ltr(doc, 'Purpose: Authenticate existing users and grant system access.')
    add_paragraph_with_ltr(doc, 'Steps:')
    para = add_paragraph_with_ltr(doc, '1. User accesses login page')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. User enters email and password')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. System validates input (non-empty fields)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. System searches for user with matching email (case-insensitive)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. If user not found, display "Invalid email or password" and return to step 2')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. If user found, retrieve stored password hash')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. Hash the entered password using SHA-256')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '8. Compare entered password hash with stored hash')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '9. If hashes do not match, display "Invalid email or password" and return to step 2')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '10. If hashes match, generate JWT token with user information')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '11. Store token in client-side storage')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '12. Return success response with token and user data')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '13. Redirect user to role-specific dashboard')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '3.2.3 CV Generation Algorithm', 3)
    add_paragraph_with_ltr(doc, 'Purpose: Automatically generate professional CV from graduate profile data.')
    add_paragraph_with_ltr(doc, 'Steps:')
    para = add_paragraph_with_ltr(doc, '1. User accesses profile page')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. User clicks "Generate CV" button')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. System retrieves graduate profile data (name, education, skills, experience, etc.)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. System retrieves user information (email, full name)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. User selects CV template (Template 1, 2, or 3)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. System combines profile data with selected template structure')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. System applies styling and formatting according to template')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '8. System calculates profile completeness percentage')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '9. System renders CV in HTML format')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '10. Display preview of generated CV')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '11. User can export CV as PDF')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '12. If PDF export requested: Convert HTML to canvas/image, Generate PDF document, Download PDF file to user\'s device')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '13. User can print CV directly from browser')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '14. Save CV generation history (optional)')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '3.2.4 Job Application Algorithm', 3)
    add_paragraph_with_ltr(doc, 'Purpose: Enable graduates to apply for job postings.')
    add_paragraph_with_ltr(doc, 'Steps:')
    para = add_paragraph_with_ltr(doc, '1. Graduate browses available jobs')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. Graduate views job details (title, description, requirements, salary)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Graduate clicks "Apply" button')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. System verifies user authentication (JWT token)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. System checks user role (must be Graduate)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. If not authenticated or wrong role, redirect to login')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. System retrieves graduate profile')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '8. If graduate profile incomplete, prompt user to complete profile')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '9. System checks if graduate already applied for this job')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '10. If already applied, display message and prevent duplicate application')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '11. If not applied, display application form')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '12. Graduate enters optional cover letter')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '13. Graduate submits application')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '14. System creates application record with: Application ID, Job ID, Graduate ID, Status: "pending", Cover letter, Application date')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '15. Save application to persistent storage')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '16. Notify company of new application (optional)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '17. Display success message to graduate')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '18. Update job application count')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '3.2.5 Job Posting Algorithm (Company)', 3)
    add_paragraph_with_ltr(doc, 'Purpose: Enable companies to post job opportunities.')
    add_paragraph_with_ltr(doc, 'Steps:')
    para = add_paragraph_with_ltr(doc, '1. Company user logs into system')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. Company accesses company portal')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Company clicks "Post New Job" button')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. System verifies authentication and company role')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. System checks if company profile exists')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. If company profile missing, auto-create basic profile')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. Display job posting form')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '8. Company fills job details: Job title, Description, Location, Salary (optional), Required skills, Employment type (full-time, part-time, contract)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '9. System validates required fields')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '10. If validation fails, display errors and return to step 8')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '11. If validation passes, create job record with: Job ID, Company ID, All job details, Status: "active", Creation date')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '12. Save job to persistent storage')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '13. Display success message')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '14. Job appears in job listings for graduates')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '15. Update company\'s job count')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '3.2.6 Application Status Update Algorithm (Company)', 3)
    add_paragraph_with_ltr(doc, 'Purpose: Allow companies to accept or reject job applications.')
    add_paragraph_with_ltr(doc, 'Steps:')
    para = add_paragraph_with_ltr(doc, '1. Company accesses company portal')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. Company views "My Jobs" section')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Company clicks "View Applications" for a specific job')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. System retrieves all applications for that job')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. System displays list of applicants with: Applicant name, Major and university, GPA, Skills, Application status, Cover letter')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. Company clicks "Accept" or "Reject" for an application')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. System verifies company owns the job')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '8. System updates application status: If Accept: status = "accepted", If Reject: status = "rejected"')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '9. Save updated application to persistent storage')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '10. Display success message to company')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '11. Update application status badge in UI')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '12. Notify graduate of status change (optional)')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '3.2.7 Profile Update Algorithm (Graduate)', 3)
    add_paragraph_with_ltr(doc, 'Purpose: Allow graduates to update their profile information.')
    add_paragraph_with_ltr(doc, 'Steps:')
    para = add_paragraph_with_ltr(doc, '1. Graduate accesses profile page')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. System loads current profile data')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Graduate clicks "Edit Profile" button')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. System displays edit form with current values pre-filled')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. Graduate modifies fields: Personal information (name, date of birth, gender), Education (university, major, GPA), Skills, Experience, Projects, Profile picture (optional)')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. System validates updated data')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. If validation fails, display errors')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '8. If validation passes, update graduate profile record')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '9. Save updated profile to persistent storage')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '10. Update profile completeness percentage')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '11. Display success message')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '12. Refresh profile display with new data')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '13. If CV was generated, prompt user to regenerate CV')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '3.3 Flowcharts', 2)
    add_paragraph_with_ltr(doc,
        'The following flowcharts illustrate the logical flow of major system processes. These diagrams were prepared '
        'by the researcher to visualize the step-by-step procedures implemented in JoinWork.')
    
    add_heading_with_ltr(doc, '3.3.1 Login Flowchart', 3)
    add_paragraph_with_ltr(doc,
        'The login process begins with user input validation, proceeds through authentication verification, and concludes '
        'with role-based redirection to the appropriate dashboard.')
    
    add_heading_with_ltr(doc, '3.3.2 Registration Flowchart', 3)
    add_paragraph_with_ltr(doc,
        'The registration process includes role selection, form validation, email uniqueness checking, password hashing, '
        'profile creation, and automatic login upon successful registration.')
    
    add_heading_with_ltr(doc, '3.3.3 Job Application Flowchart', 3)
    add_paragraph_with_ltr(doc,
        'The job application process verifies user authentication and role, checks profile completeness, prevents duplicate '
        'applications, and creates application records with pending status.')
    
    add_heading_with_ltr(doc, '3.3.4 CV Generation Flowchart', 3)
    add_paragraph_with_ltr(doc,
        'The CV generation process retrieves user data, allows template selection, combines data with template structure, '
        'renders the CV, and provides export and print options.')
    
    add_heading_with_ltr(doc, '3.4 Interface Explanation', 2)
    
    add_heading_with_ltr(doc, '3.4.1 Graduate Interface', 3)
    
    add_heading_with_ltr(doc, 'Dashboard Page:', 4)
    para = add_paragraph_with_ltr(doc, '• Overview of available jobs')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Profile completeness indicator')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Quick access to profile, jobs, and workshops')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Statistics (applications sent, profile views)')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, 'Profile Page:', 4)
    para = add_paragraph_with_ltr(doc, '• Personal information display and editing')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Education details (university, major, GPA)')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Skills and experience')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• CV generation section with template selection')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Profile picture upload')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Export/print CV functionality')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, 'Jobs Page:', 4)
    para = add_paragraph_with_ltr(doc, '• List of available job postings')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Filtering options (location, salary, skills)')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Job details view')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Apply button for each job')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Search functionality')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, 'Workshops Page:', 4)
    para = add_paragraph_with_ltr(doc, '• List of available training workshops')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Workshop details and registration')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Category filtering')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '3.4.2 Company Interface', 3)
    
    add_heading_with_ltr(doc, 'Company Portal:', 4)
    para = add_paragraph_with_ltr(doc, '• Dashboard with job statistics')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• "My Jobs" section listing all posted jobs')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• "Post New Job" functionality')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Application management: View applications per job, Accept/reject applications, View applicant profiles')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Company profile management')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, 'Job Posting Form:', 4)
    para = add_paragraph_with_ltr(doc, '• Job title, description, location')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Salary and employment type')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Required skills')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Job status management')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '3.4.3 Ministry Interface', 3)
    
    add_heading_with_ltr(doc, 'Ministry Portal:', 4)
    para = add_paragraph_with_ltr(doc, '• Analytics dashboard')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Graduate statistics')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Employment trends')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Workshop management')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Reports generation')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Announcements management')
    para.style = 'List Bullet'
    
    doc.add_page_break()
    
    # CHAPTER 4: IMPLEMENTATION AND RESULTS
    add_heading_with_ltr(doc, 'CHAPTER 4: IMPLEMENTATION AND RESULTS', 1)
    
    add_heading_with_ltr(doc, '4.1 Chapter Introduction', 2)
    add_paragraph_with_ltr(doc,
        'This chapter demonstrates the practical implementation of JoinWork, showing how the system functions in real-world '
        'scenarios. It includes interface descriptions, system walkthroughs, and discussion of results and achievements.')
    
    add_heading_with_ltr(doc, '4.2 Practical System Demonstration', 2)
    
    add_heading_with_ltr(doc, '4.2.1 System Overview', 3)
    add_paragraph_with_ltr(doc,
        'JoinWork operates as a web-based platform accessible through standard web browsers. The system follows a '
        'client-server architecture where the frontend (HTML, CSS, JavaScript) communicates with the backend (Python Flask) '
        'through RESTful API endpoints.')
    
    add_heading_with_ltr(doc, '4.2.2 Usage Sequence', 3)
    
    add_heading_with_ltr(doc, 'For Graduates:', 4)
    para = add_paragraph_with_ltr(doc, '1. Registration → Create account with graduate role')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. Profile Completion → Fill in personal, educational, and professional details')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. CV Generation → Select template and generate professional CV')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. Job Browsing → Search and filter available job opportunities')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. Job Application → Apply for jobs with optional cover letter')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. Application Tracking → Monitor application status')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, 'For Companies:', 4)
    para = add_paragraph_with_ltr(doc, '1. Registration → Create company account')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. Company Profile → Complete company information')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Job Posting → Create and publish job opportunities')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. Application Review → View and manage received applications')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. Decision Making → Accept or reject applications')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. Applicant Evaluation → View detailed graduate profiles')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, 'For Ministry:', 4)
    para = add_paragraph_with_ltr(doc, '1. Access Portal → Login with ministry credentials')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. View Analytics → Monitor employment statistics')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Manage Workshops → Create and manage training programs')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. Generate Reports → Export data and insights')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '4.3 Interface Descriptions', 2)
    
    add_heading_with_ltr(doc, '4.3.1 Login Page', 3)
    add_paragraph_with_ltr(doc,
        'The login interface provides email and password input fields, role-based redirection after authentication, '
        'bilingual support (English/Arabic), dark mode toggle, and error handling with clear messages.')
    add_paragraph_with_ltr(doc, 'Key Features: Input validation, Secure authentication, Responsive design, User-friendly interface')
    
    add_heading_with_ltr(doc, '4.3.2 Registration Page', 3)
    add_paragraph_with_ltr(doc,
        'The registration interface includes role selection (Graduate, Company, Ministry), dynamic form fields based on '
        'selected role, gender selection and date of birth calendar for graduates, optional profile picture upload, '
        'form validation, and bilingual labels and placeholders.')
    add_paragraph_with_ltr(doc,
        'Key Features: Conditional field display, Date picker for birth date, Image upload with preview, Real-time validation')
    
    add_heading_with_ltr(doc, '4.3.3 Profile Page', 3)
    add_paragraph_with_ltr(doc,
        'The graduate profile page offers personal information section, education details (university, major, GPA), '
        'skills and experience, projects portfolio, profile picture display, CV generation section with multiple templates, '
        'export to PDF functionality, print option, and profile completeness indicator.')
    add_paragraph_with_ltr(doc,
        'Key Features: Comprehensive data display, Easy editing capability, Professional CV generation, Multiple export options')
    
    add_heading_with_ltr(doc, '4.3.4 Job Listings Page', 3)
    add_paragraph_with_ltr(doc,
        'The jobs page displays grid/list view of available jobs, job cards with key information (title, company, location, '
        'salary), filtering options (location, salary range, skills), search functionality, apply button for each job, '
        'and job details modal.')
    add_paragraph_with_ltr(doc,
        'Key Features: Efficient job browsing, Advanced filtering, Quick application process, Responsive layout')
    
    add_heading_with_ltr(doc, '4.3.5 CV Builder', 3)
    add_paragraph_with_ltr(doc,
        'The CV generation feature provides three professional templates, automatic data population from profile, real-time '
        'preview, color contrast optimization for readability, PDF export with proper formatting, and print-ready output.')
    add_paragraph_with_ltr(doc,
        'Key Features: Multiple template options, Professional appearance, Easy customization, High-quality output')
    
    add_heading_with_ltr(doc, '4.3.6 Company Portal', 3)
    add_paragraph_with_ltr(doc,
        'The company interface includes dashboard with statistics, job management section, application review interface, '
        'applicant profile viewing, accept/reject functionality, and company profile management.')
    add_paragraph_with_ltr(doc,
        'Key Features: Comprehensive job management, Efficient application handling, Detailed applicant information, Status tracking')
    
    add_heading_with_ltr(doc, '4.4 Results Discussion', 2)
    
    add_heading_with_ltr(doc, '4.4.1 Objective Achievement', 3)
    
    add_paragraph_with_ltr(doc, 'Objective 1: Create a Digital Platform')
    add_paragraph_with_ltr(doc, '✅ Achieved: JoinWork successfully provides a web-based platform connecting graduates with employers. '
        'The system is fully functional with all core features implemented.')
    
    add_paragraph_with_ltr(doc, 'Objective 2: Facilitate CV Creation')
    add_paragraph_with_ltr(doc, '✅ Achieved: The automated CV generation system produces professional CVs from profile data. '
        'Multiple templates offer variety, and PDF export ensures compatibility.')
    
    add_paragraph_with_ltr(doc, 'Objective 3: Enable Job Posting')
    add_paragraph_with_ltr(doc, '✅ Achieved: Companies can post jobs with detailed requirements. The job management system '
        'allows full CRUD operations.')
    
    add_paragraph_with_ltr(doc, 'Objective 4: Streamline Application Process')
    add_paragraph_with_ltr(doc, '✅ Achieved: The application process is simplified to a few clicks. Companies can efficiently '
        'manage and review applications.')
    
    add_paragraph_with_ltr(doc, 'Objective 5: Reduce the Gap')
    add_paragraph_with_ltr(doc, '✅ Achieved: The platform bridges academia and industry by providing relevant job opportunities '
        'and facilitating direct communication between graduates and employers.')
    
    add_paragraph_with_ltr(doc, 'Objective 6: Support Recent Graduates')
    add_paragraph_with_ltr(doc, '✅ Achieved: The system is designed with recent graduates in mind, featuring simplified '
        'interfaces and comprehensive guidance.')
    
    add_paragraph_with_ltr(doc, 'Objective 7: Provide Analytics')
    add_paragraph_with_ltr(doc, '✅ Achieved: The Ministry portal provides insights into employment trends and graduate statistics.')
    
    add_heading_with_ltr(doc, '4.4.2 System Usability', 3)
    
    add_paragraph_with_ltr(doc, 'Ease of Use:')
    para = add_paragraph_with_ltr(doc, '• Intuitive interface design')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Clear navigation')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Bilingual support for accessibility')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Responsive design for various devices')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Minimal learning curve')
    para.style = 'List Bullet'
    
    add_paragraph_with_ltr(doc, 'User Experience:')
    para = add_paragraph_with_ltr(doc, '• Fast response times')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Clear error messages')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Success confirmations')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Helpful tooltips and guidance')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Professional appearance')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '4.4.3 System Performance', 3)
    
    add_paragraph_with_ltr(doc, 'Response Time:')
    para = add_paragraph_with_ltr(doc, '• Fast page loading')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Efficient API responses')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Quick data retrieval')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Smooth user interactions')
    para.style = 'List Bullet'
    
    add_paragraph_with_ltr(doc, 'Reliability:')
    para = add_paragraph_with_ltr(doc, '• Persistent data storage')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Error handling mechanisms')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Input validation')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Secure authentication')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '4.4.4 Technical Achievements', 3)
    para = add_paragraph_with_ltr(doc, '1. Full-Stack Development: Successfully implemented both frontend and backend components')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. RESTful API: Well-structured API following REST principles')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Security: Implemented JWT authentication and password hashing')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. Data Persistence: JSON-based storage system for data reliability')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. Bilingual Support: Complete Arabic and English interface')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. CV Generation: Automated professional CV creation with multiple templates')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '7. Role-Based Access: Proper authorization for different user types')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '4.4.5 Limitations and Challenges', 3)
    
    add_paragraph_with_ltr(doc, 'Current Limitations:')
    para = add_paragraph_with_ltr(doc, '1. Scalability: JSON file storage may become inefficient with very large datasets')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. Real-time Features: No real-time notifications or updates')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Advanced Search: Limited to basic filtering; no advanced search algorithms')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. Mobile App: Web-only; no native mobile application')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. Email Integration: No email notifications for applications or status updates')
    para.style = 'List Number'
    
    add_paragraph_with_ltr(doc, 'Challenges Overcome:')
    para = add_paragraph_with_ltr(doc, '1. Data Persistence: Implemented JSON storage to maintain data across server restarts')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. Authentication: Successfully integrated JWT for secure user sessions')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. CV Generation: Developed multiple templates with proper formatting')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. Bilingual Support: Implemented comprehensive Arabic/English interface')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. Role Management: Properly handled different user roles and permissions')
    para.style = 'List Number'
    
    doc.add_page_break()
    
    # CHAPTER 5: CONCLUSIONS AND FUTURE WORK
    add_heading_with_ltr(doc, 'CHAPTER 5: CONCLUSIONS AND FUTURE WORK', 1)
    
    add_heading_with_ltr(doc, '5.1 Conclusions', 2)
    
    add_heading_with_ltr(doc, '5.1.1 Project Accomplishments', 3)
    add_paragraph_with_ltr(doc,
        'JoinWork successfully addresses the problem of connecting graduates with employment opportunities through a '
        'centralized digital platform. The system provides:')
    para = add_paragraph_with_ltr(doc, '1. Comprehensive Solution: A complete platform serving graduates, companies, and the Ministry')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '2. User-Friendly Interface: Intuitive design accessible to users with varying technical skills')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '3. Professional Tools: Automated CV generation producing high-quality documents')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '4. Efficient Processes: Streamlined job posting and application management')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '5. Secure System: Proper authentication and authorization mechanisms')
    para.style = 'List Number'
    para = add_paragraph_with_ltr(doc, '6. Persistent Data: Reliable data storage ensuring information retention')
    para.style = 'List Number'
    
    add_heading_with_ltr(doc, '5.1.2 Problem Resolution', 3)
    add_paragraph_with_ltr(doc,
        'The research problem—the absence of a unified platform connecting graduates with the job market—has been addressed through:')
    para = add_paragraph_with_ltr(doc, '• Centralized Platform: All job opportunities and graduate profiles in one system')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Automated CV Creation: Professional CVs generated automatically')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Organized Database: Systematic storage and retrieval of graduate and job data')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Verified Information: User authentication ensures data reliability')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Bridging the Gap: Direct connection between academic achievements and job requirements')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '5.1.3 Scientific and Practical Value', 3)
    
    add_paragraph_with_ltr(doc, 'Scientific Value:')
    para = add_paragraph_with_ltr(doc, '• Demonstrates application of web technologies in solving real-world problems')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Implements modern software engineering practices')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Showcases full-stack development capabilities')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Applies security best practices in web applications')
    para.style = 'List Bullet'
    
    add_paragraph_with_ltr(doc, 'Practical Value:')
    para = add_paragraph_with_ltr(doc, '• Provides immediate benefit to graduates seeking employment')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Assists companies in finding qualified candidates')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Supports Ministry in monitoring employment trends')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Contributes to digital transformation in employment services')
    para.style = 'List Bullet'
    para = add_paragraph_with_ltr(doc, '• Free and accessible platform for all users')
    para.style = 'List Bullet'
    
    add_heading_with_ltr(doc, '5.2 Objective Fulfillment', 2)
    
    add_paragraph_with_ltr(doc, 'Objective 1: Create a Digital Platform')
    add_paragraph_with_ltr(doc, 'Status: ✅ Fully Achieved - Web-based platform successfully developed with all core features implemented and functional.')
    
    add_paragraph_with_ltr(doc, 'Objective 2: Facilitate CV Creation')
    add_paragraph_with_ltr(doc, 'Status: ✅ Fully Achieved - Automated CV generation system operational with multiple professional templates and PDF export.')
    
    add_paragraph_with_ltr(doc, 'Objective 3: Enable Job Posting')
    add_paragraph_with_ltr(doc, 'Status: ✅ Fully Achieved - Companies can post jobs with full details and manage them through comprehensive job management system.')
    
    add_paragraph_with_ltr(doc, 'Objective 4: Streamline Application Process')
    add_paragraph_with_ltr(doc, 'Status: ✅ Fully Achieved - Simple application process for graduates and efficient application management for companies.')
    
    add_paragraph_with_ltr(doc, 'Objective 5: Reduce the Gap')
    add_paragraph_with_ltr(doc, 'Status: ✅ Achieved - Platform bridges academia and industry through relevant job opportunities and direct communication.')
    
    add_paragraph_with_ltr(doc, 'Objective 6: Support Recent Graduates')
    add_paragraph_with_ltr(doc, 'Status: ✅ Achieved - System designed with recent graduates in mind, featuring simplified interfaces and comprehensive guidance.')
    
    add_paragraph_with_ltr(doc, 'Objective 7: Provide Analytics')
    add_paragraph_with_ltr(doc, 'Status: ✅ Achieved - Ministry portal provides insights into employment trends and graduate statistics.')
    
    add_heading_with_ltr(doc, '5.3 Future Work', 2)
    
    add_heading_with_ltr(doc, '5.3.1 Mobile Application Development', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Develop native mobile applications for iOS and Android platforms to increase accessibility '
        'and convenience for users on mobile devices.')
    add_paragraph_with_ltr(doc, 'Benefits: Push notifications for job matches and application updates, Offline functionality '
        'for profile viewing, Camera integration for profile pictures, Location-based job recommendations')
    
    add_heading_with_ltr(doc, '5.3.2 Job Recommendation System', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Implement an AI-powered recommendation system that suggests relevant jobs to graduates based '
        'on their profile, skills, and preferences.')
    add_paragraph_with_ltr(doc, 'Features: Machine learning algorithms for job matching, Personalized job recommendations, '
        'Skill gap analysis, Career path suggestions')
    
    add_heading_with_ltr(doc, '5.3.3 Multi-University Integration', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Extend the platform to integrate with multiple universities, allowing automatic profile '
        'creation for graduates and direct data synchronization.')
    add_paragraph_with_ltr(doc, 'Benefits: Automated graduate registration, Verified academic credentials, University-specific '
        'job postings, Alumni network features')
    
    add_heading_with_ltr(doc, '5.3.4 Advanced CV Analysis with AI', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Implement artificial intelligence to analyze CVs, provide improvement suggestions, and optimize '
        'content for better job matching.')
    add_paragraph_with_ltr(doc, 'Features: CV quality scoring, Keyword optimization suggestions, ATS (Applicant Tracking System) '
        'compatibility check, Industry-specific recommendations')
    
    add_heading_with_ltr(doc, '5.3.5 Real-Time Communication', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Add real-time messaging system between graduates and companies for direct communication and '
        'interview scheduling.')
    add_paragraph_with_ltr(doc, 'Features: In-platform messaging, Video interview scheduling, Notification system, Interview '
        'preparation resources')
    
    add_heading_with_ltr(doc, '5.3.6 Database Migration', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Migrate from JSON file storage to a relational database (MySQL or PostgreSQL) for better '
        'scalability and performance.')
    add_paragraph_with_ltr(doc, 'Benefits: Improved query performance, Better data relationships, Enhanced data integrity, '
        'Support for complex queries')
    
    add_heading_with_ltr(doc, '5.3.7 Advanced Analytics Dashboard', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Develop comprehensive analytics dashboards with data visualization, trend analysis, and '
        'predictive insights.')
    add_paragraph_with_ltr(doc, 'Features: Interactive charts and graphs, Employment trend predictions, Skill demand analysis, '
        'Regional employment statistics, Export capabilities for reports')
    
    add_heading_with_ltr(doc, '5.3.8 Integration with Social Media', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Allow users to import profile information from LinkedIn or other professional networks and share '
        'job postings on social media.')
    add_paragraph_with_ltr(doc, 'Features: Social media login options, Profile import functionality, Social sharing capabilities, '
        'Network expansion tools')
    
    add_heading_with_ltr(doc, '5.3.9 Workshop and Training Management', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Expand the workshop system to include online courses, certification programs, and skill assessment tests.')
    add_paragraph_with_ltr(doc, 'Features: Online course platform, Certification tracking, Skill assessment quizzes, Progress monitoring')
    
    add_heading_with_ltr(doc, '5.3.10 Multi-Language Support', 3)
    add_paragraph_with_ltr(doc,
        'Proposed Enhancement: Extend language support beyond Arabic and English to include Kurdish and other regional languages.')
    add_paragraph_with_ltr(doc, 'Benefits: Broader user base, Regional accessibility, Cultural inclusivity')
    
    doc.add_page_break()
    
    # REFERENCES
    add_heading_with_ltr(doc, 'REFERENCES', 1)
    
    refs = [
        '[1] Berners-Lee, T. (1991). "HTML: A Markup Language for the World Wide Web." World Wide Web Consortium.',
        '[2] Wium Lie, H. (1994). "Cascading Style Sheets: A Proposal." World Wide Web Consortium.',
        '[3] Eich, B. (1995). "JavaScript: The Definitive Guide." O\'Reilly Media.',
        '[4] van Rossum, G. (1991). "Python Programming Language." Python Software Foundation.',
        '[5] Ronacher, A. (2010). "Flask Web Framework." Flask Documentation. Available: https://flask.palletsprojects.com/',
        '[6] Crockford, D. (2001). "JSON: JavaScript Object Notation." RFC 4627, IETF.',
        '[7] Jones, M., Bradley, J., & Sakimura, N. (2015). "JSON Web Token (JWT)." RFC 7519, IETF.',
        '[8] National Institute of Standards and Technology. (2001). "Secure Hash Standard (SHS)." FIPS PUB 180-4.',
        '[9] Fielding, R. T. (2000). "Architectural Styles and the Design of Network-based Software Architectures." University of California, Irvine.',
        '[10] W3C. (2021). "Web Content Accessibility Guidelines (WCAG) 2.1." World Wide Web Consortium.',
        '[11] MDN Web Docs. (2023). "JavaScript Guide." Mozilla Developer Network. Available: https://developer.mozilla.org/en-US/docs/Web/JavaScript/Guide',
        '[12] Flask Documentation. (2023). "Flask: Web Development, One Drop at a Time." Available: https://flask.palletsprojects.com/',
        '[13] Python Software Foundation. (2023). "Python Documentation." Available: https://www.python.org/doc/',
        '[14] W3Schools. (2023). "HTML Tutorial." Available: https://www.w3schools.com/html/',
        '[15] W3Schools. (2023). "CSS Tutorial." Available: https://www.w3schools.com/css/',
        '[16] LinkedIn Corporation. (2023). "LinkedIn: About Us." Available: https://about.linkedin.com/',
        '[17] Bayt.com. (2023). "Bayt.com: The Middle East\'s #1 Job Site." Available: https://www.bayt.com/',
        '[18] Mozilla Developer Network. (2023). "Cross-Origin Resource Sharing (CORS)." Available: https://developer.mozilla.org/en-US/docs/Web/HTTP/CORS',
        '[19] jsPDF Documentation. (2023). "jsPDF: A library to generate PDFs in JavaScript." Available: https://github.com/parallax/jsPDF',
        '[20] html2canvas Documentation. (2023). "html2canvas: Screenshots with JavaScript." Available: https://html2canvas.hertzen.com/'
    ]
    
    for ref in refs:
        para = add_paragraph_with_ltr(doc, ref)
        para.style = 'List Bullet'
    
    # Save document
    import os
    # Get the directory where this script is located
    script_dir = os.path.dirname(os.path.abspath(__file__)) if '__file__' in globals() else os.getcwd()
    filename = os.path.join(script_dir, 'Graduation_Project_Ameer.docx')
    doc.save(filename)
    print(f'\n✅ Thesis document created successfully: {filename}')
    print(f'📄 File location: {os.path.abspath(filename)}\n')

if __name__ == '__main__':
    import os
    import sys
    # Ensure we're in the right directory
    script_path = os.path.abspath(__file__)
    script_dir = os.path.dirname(script_path)
    os.chdir(script_dir)
    create_thesis_document()

